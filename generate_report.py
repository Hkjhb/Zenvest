"""
Generates a Word document report for the Zenvest project.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────────────────────────────────────
# Page margins
# ─────────────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.54)

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
TITLE_COLOR  = RGBColor(0x1A, 0x56, 0x76)   # dark teal
H1_COLOR     = RGBColor(0x1A, 0x56, 0x76)
H2_COLOR     = RGBColor(0x2E, 0x74, 0x96)
ACCENT_COLOR = RGBColor(0x0D, 0x47, 0x6E)
TABLE_HEADER  = RGBColor(0x1A, 0x56, 0x76)
TABLE_ALT     = RGBColor(0xEA, 0xF4, 0xFB)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",  "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",   "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color","auto"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def heading(text, level=1, color=None, bold=True, center=False, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    size_map = {1: 16, 2: 14, 3: 12, 4: 11}
    run.font.size = Pt(size_map.get(level, 12))
    if color:
        run.font.color.rgb = color
    elif level == 1:
        run.font.color.rgb = H1_COLOR
    elif level == 2:
        run.font.color.rgb = H2_COLOR
    return p


def body(text, indent=0, bold=False, italic=False, size=11, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(0.75 + level * 0.5)
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.space_before  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def page_break():
    doc.add_page_break()


def horizontal_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A5676")
    pBdr.append(bot)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# 1.  COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
r = p.add_run("ZENVEST")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = TITLE_COLOR

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(4)
r2 = p2.add_run("A Full-Stack MERN Paper Trading Simulator")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
r2.italic = True

horizontal_rule()

for label, value in [
    ("Subject",     "Full Stack Development (FSD)"),
    ("Department",  "Department of Computer Science & Engineering / Information Technology"),
    ("College",     "[College Name]"),
    ("Academic Year", "2025 – 2026"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"{label}:  ")
    r.bold = True
    r.font.size = Pt(12)
    r2 = p.add_run(value)
    r2.font.size = Pt(12)

horizontal_rule()

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Team Members")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = TITLE_COLOR

for member in [
    ("Harsh Kumar",  "[Roll No.]", "[Branch]"),
    ("[Member 2]",   "[Roll No.]", "[Branch]"),
    ("[Member 3]",   "[Roll No.]", "[Branch]"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"{member[0]}   |   {member[1]}   |   {member[2]}")
    r.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Guided By:  [Faculty Guide Name]")
r.font.size = Pt(11)
r.italic = True

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 2.  ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────
heading("2. Abstract", level=1)
horizontal_rule()

body(
    "Zenvest is a full-stack web application that simulates an equity stock trading "
    "platform, modelled on the operational design of platforms such as Zerodha. Built "
    "on the MERN stack (MongoDB, Express.js, React, Node.js), the system provides users "
    "with a realistic, risk-free environment to learn and practise stock trading using "
    "virtual capital."
)
body(
    "The application comprises three independently deployable components: a React-based "
    "marketing landing site, a React-based trading dashboard, and a Node.js/Express REST "
    "API backed by MongoDB. Real-time NSE stock prices for fourteen major Indian equities "
    "are sourced through the Yahoo Finance API and cached server-side to minimise external "
    "calls. Users can place market orders (executed instantly at the live price) and limit "
    "orders (queued and automatically executed when the market price meets the target). "
    "Portfolio analytics—including P&L, current value, day change, and allocation "
    "visualisations—are presented through Chart.js-powered charts."
)
body(
    "Authentication is implemented with PBKDF2-SHA-512 password hashing and stateless "
    "bearer-token sessions. A dual trading-mode toggle (Paper / Live) allows users to "
    "experience a simulated payment gateway for fund deposits while maintaining a separate "
    "virtual balance for practice. The project demonstrates proficiency in component-based "
    "UI architecture, RESTful API design, real-time data integration, and secure "
    "authentication patterns."
)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 3.  TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
heading("3. Table of Contents", level=1)
horizontal_rule()

toc_items = [
    ("1.", "Cover Page",                              "1"),
    ("2.", "Abstract",                                "2"),
    ("3.", "Table of Contents",                       "3"),
    ("4.", "Introduction",                            "4"),
    ("   4.1", "Problem Statement",                   "4"),
    ("   4.2", "Objectives of the Project",           "4"),
    ("   4.3", "Scope of the Application",            "5"),
    ("5.", "Technology Stack",                        "5"),
    ("   5.1", "Frontend – React",                   "5"),
    ("   5.2", "Backend – Node.js & Express",         "6"),
    ("   5.3", "Database – MongoDB",                  "6"),
    ("   5.4", "Tools Used",                          "6"),
    ("6.", "System Design",                           "7"),
    ("   6.1", "Architecture Diagram",                "7"),
    ("   6.2", "ER Diagram / Database Schema",        "7"),
    ("7.", "Module Description",                      "8"),
    ("8.", "Implementation Details",                  "10"),
    ("   8.1", "Overview of Features",                "10"),
    ("   8.2", "Code Structure",                      "11"),
    ("   8.3", "Key Logic Snippets",                  "12"),
    ("9.", "Deployment Details",                      "13"),
    ("10.", "Project Implementation Status Table",    "14"),
    ("11.", "References",                             "15"),
    ("12.", "Appendix – Screenshots",                 "16"),
]

tbl = doc.add_table(rows=len(toc_items), cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
col_widths = [Cm(1.5), Cm(11.0), Cm(1.5)]
for i, row in enumerate(tbl.rows):
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, val="none", color="FFFFFF")
    num, title, pg = toc_items[i]
    tbl.rows[i].cells[0].paragraphs[0].add_run(num).font.size = Pt(11)
    r = tbl.rows[i].cells[1].paragraphs[0].add_run(title)
    r.font.size = Pt(11)
    tbl.rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tbl.rows[i].cells[2].paragraphs[0].add_run(pg).font.size = Pt(11)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 4.  INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
heading("4. Introduction", level=1)
horizontal_rule()

heading("4.1  Problem Statement", level=2)
body(
    "Stock trading is an important financial skill, yet the barrier to entry remains high "
    "for beginners. Real capital is required to practise on live markets, and the risk of "
    "financial loss discourages new participants from experimenting with trading strategies. "
    "Existing educational platforms are either overly simplified—lacking real price data—or "
    "complex enough to overwhelm novices."
)
body(
    "There is a clear need for a platform that bridges the gap: one that delivers a "
    "professional-grade trading experience using live market data, while eliminating the "
    "financial risk through virtual capital. Zenvest addresses this need by providing an "
    "end-to-end simulated trading environment that mirrors the workflow of a real-world "
    "equity trading platform."
)

heading("4.2  Objectives of the Project", level=2)
for obj in [
    "Develop a responsive, multi-page marketing website that communicates the platform's value proposition to prospective users.",
    "Build a secure user authentication system using PBKDF2-SHA-512 password hashing and stateless bearer-token sessions.",
    "Implement a real-time stock watchlist that fetches live NSE prices for fourteen major Indian equities via the Yahoo Finance API.",
    "Enable market order execution (instant, at live price) and limit order execution (background job, price-triggered) for buy and sell actions.",
    "Provide portfolio analytics including P&L calculation, holdings visualisation, and order history.",
    "Simulate a payment gateway for fund deposits under a 'Live' trading mode while maintaining a separate virtual balance for practice.",
    "Deploy the application across separate frontend, dashboard, and backend services using industry-standard tooling.",
    "Demonstrate a clean separation of concerns through a three-tier MERN architecture.",
]:
    bullet(obj)

heading("4.3  Scope of the Application", level=2)
body(
    "The application is scoped as an educational trading simulator and does not connect "
    "to any real brokerage or execute transactions on a live exchange. The scope includes:"
)
for item in [
    "Fourteen NSE-listed equities (INFY, TCS, RELIANCE, HDFCBANK, WIPRO, ONGC, M&M, BHARTIARTL, HINDUNILVR, ITC, SBIN, TATAPOWER, KPITTECH, QUICKHEAL) and two market indices (NIFTY 50, SENSEX).",
    "Virtual paper-trading with an initial balance of ₹50,000 and the ability to add up to ₹1,00,000 in additional virtual funds per transaction.",
    "A simulated 'Live' mode with a separate real balance, funded through a mock payment gateway (Card / UPI / Net Banking) with a ceiling of ₹5,00,000 per deposit.",
    "Background limit-order matching engine running every 30 seconds.",
    "7-day simulated price trend charts, doughnut allocation charts, and bar charts for portfolio analytics.",
    "Portfolio reset functionality to restore the account to its initial seeded state.",
]:
    bullet(item)

body(
    "Out of scope: real-money transactions, derivatives (F&O), mutual funds, IPOs, "
    "tax reporting, and integration with SEBI-regulated brokerage APIs."
)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 5.  TECHNOLOGY STACK
# ─────────────────────────────────────────────────────────────────────────────
heading("5. Technology Stack", level=1)
horizontal_rule()

heading("5.1  Frontend – React", level=2)
body(
    "Both the marketing landing site and the trading dashboard are built with React 18. "
    "React's component-based architecture enables reusable UI elements, while React Router "
    "DOM provides client-side navigation without full page reloads."
)
for item in [
    "React 18 – declarative UI with hooks (useState, useEffect, useContext, useRef)",
    "React Router DOM v6 – client-side routing with nested routes and protected paths",
    "Chart.js / react-chartjs-2 – Doughnut, Bar, and Line charts for portfolio analytics",
    "Material-UI (MUI) – Tooltip, Grow transition, and icon components",
    "Bootstrap 5 – responsive grid and utility classes for the landing pages",
    "Axios – HTTP client with request interceptors for automatic bearer-token attachment",
    "Custom CSS – component-level stylesheets for the trading dashboard",
]:
    bullet(item)

heading("5.2  Backend – Node.js & Express", level=2)
body(
    "The backend is a REST API server built with Node.js and Express 4.18. It handles "
    "authentication, portfolio management, live price serving, and asynchronous order matching."
)
for item in [
    "Node.js – non-blocking I/O runtime for handling concurrent requests and background jobs",
    "Express 4.18 – minimal, unopinionated web framework for defining REST routes",
    "Mongoose 8.2 – MongoDB object modelling with schema validation and relationship references",
    "yahoo-finance2 3.14 – unofficial Yahoo Finance SDK for fetching live NSE stock quotes",
    "cors 2.8 – Cross-Origin Resource Sharing middleware for frontend-backend communication",
    "body-parser 1.20 – JSON request body parsing",
    "crypto (Node built-in) – PBKDF2 key derivation and random token generation",
]:
    bullet(item)

heading("5.3  Database – MongoDB", level=2)
body(
    "MongoDB is used as the primary data store. Its document-oriented model maps naturally "
    "to the varying structures of holdings, positions, and orders. Mongoose schemas enforce "
    "data integrity at the application layer."
)
for item in [
    "Four collections: users, holdings, positions, orders",
    "ObjectId references link holdings, positions, and orders to their owning user",
    "Mongoose timestamps option on the Orders model provides automatic createdAt / updatedAt fields",
    "Indexes on userId foreign keys for efficient per-user queries",
]:
    bullet(item)

heading("5.4  Tools Used", level=2)
tools_data = [
    ("Git / GitHub",  "Version control and source-code hosting"),
    ("VS Code",       "Primary development IDE"),
    ("npm",           "Package management for all three sub-projects"),
    ("Postman",       "API endpoint testing during development"),
    ("MongoDB Atlas / Local", "Database hosting (configurable via MONGO_URL env var)"),
    ("Netlify",       "Recommended static frontend deployment"),
    ("Render / Railway", "Recommended Node.js backend deployment"),
    ("Chart.js DevTools", "Chart configuration and debugging"),
    ("Yahoo Finance", "Live NSE stock price data source"),
]
tbl = doc.add_table(rows=len(tools_data)+1, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = tbl.rows[0].cells
hdr_cells[0].paragraphs[0].add_run("Tool / Service").bold = True
hdr_cells[1].paragraphs[0].add_run("Purpose").bold = True
for cell in hdr_cells:
    set_cell_bg(cell, "1A5676")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cell.paragraphs[0].runs[0].font.size = Pt(11)
for i, (tool, purpose) in enumerate(tools_data):
    row = tbl.rows[i+1]
    row.cells[0].paragraphs[0].add_run(tool).font.size = Pt(11)
    row.cells[1].paragraphs[0].add_run(purpose).font.size = Pt(11)
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, "EAF4FB")

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 6.  SYSTEM DESIGN
# ─────────────────────────────────────────────────────────────────────────────
heading("6. System Design", level=1)
horizontal_rule()

heading("6.1  Architecture Diagram (MERN Stack)", level=2)
body(
    "Zenvest follows a standard three-tier MERN architecture. The diagram below "
    "illustrates the interaction between the three application layers and external services."
)

# ASCII-style architecture box rendered as a table
arch_rows = [
    ["USER (Browser)", "", ""],
    ["", "↑↓ HTTP/HTTPS", ""],
    ["Frontend (React – Port 3000)", "Dashboard (React – Port 3001)", ""],
    ["Landing Pages", "Trading UI + Charts", ""],
    ["", "↑↓ REST API (Axios, Bearer Token)", ""],
    ["", "Backend API (Node.js / Express – Port 3002)", ""],
    ["", "Auth | Orders | Portfolio | Live Prices", ""],
    ["", "↑↓ Mongoose ODM", "↑↓ yahoo-finance2"],
    ["", "MongoDB", "Yahoo Finance API (NSE)"],
]

body(
    "[Refer to the architecture diagram text representation below. "
    "Replace with a drawn diagram before submission.]",
    italic=True
)

arch_tbl = doc.add_table(rows=1, cols=3)
arch_tbl.style = "Table Grid"
arch_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = arch_tbl.rows[0].cells
cells[0].merge(cells[2])
p = cells[0].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

arch_text = (
    "USER (Browser)\n"
    "        |\n"
    "        | HTTP / HTTPS\n"
    "        |\n"
    "┌────────────────────┐       ┌───────────────────────────┐\n"
    "│  Frontend (React)  │       │   Dashboard (React)        │\n"
    "│  Port 3000         │       │   Port 3001                │\n"
    "│  Marketing/Landing │       │   Trading UI + Charts      │\n"
    "└────────────────────┘       └──────────┬────────────────┘\n"
    "                                        │ REST API (Bearer Token)\n"
    "                             ┌──────────┴────────────────┐\n"
    "                             │   Node.js / Express API   │\n"
    "                             │   Port 3002               │\n"
    "                             │   Auth | Orders | Prices  │\n"
    "                             └───────┬───────────┬───────┘\n"
    "                                     │           │\n"
    "                              Mongoose       yahoo-finance2\n"
    "                                     │           │\n"
    "                              ┌──────┴──┐  ┌────┴──────────┐\n"
    "                              │ MongoDB │  │ Yahoo Finance │\n"
    "                              └─────────┘  └───────────────┘\n"
)
r = p.add_run(arch_text)
r.font.name = "Courier New"
r.font.size = Pt(8)

doc.add_paragraph()

heading("6.2  ER Diagram / Database Schema", level=2)
body(
    "The database contains four collections. The users collection is the root entity; "
    "holdings, positions, and orders each hold a userId foreign key referencing the user "
    "document."
)

schema_tbl_data = [
    ("Collection", "Field",          "Type",           "Constraints / Notes"),
    ("users",      "username",       "String",         "Required, unique, trimmed"),
    ("",           "passwordHash",   "String",         "PBKDF2-SHA-512 hash, required"),
    ("",           "salt",           "String",         "16-byte random hex, required"),
    ("",           "authToken",      "String",         "Nullable; refreshed on login"),
    ("",           "openingBalance", "Number",         "Default: 50 000"),
    ("",           "cashBalance",    "Number",         "Default: 50 000"),
    ("",           "realBalance",    "Number",         "Default: 0"),
    ("holdings",   "userId",         "ObjectId → user","Required"),
    ("",           "name",           "String",         "Stock symbol e.g. INFY"),
    ("",           "qty",            "Number",         "Quantity held"),
    ("",           "avg",            "Number",         "Average buy price"),
    ("",           "price",          "Number",         "Latest traded price"),
    ("",           "net / day",      "String",         "% change strings"),
    ("positions",  "userId",         "ObjectId → user","Required"),
    ("",           "product",        "String",         "e.g. CNC"),
    ("",           "name",           "String",         "Stock symbol"),
    ("",           "qty / avg / price","Number",       "Same as holdings"),
    ("",           "isLoss",         "Boolean",        "True if position is in loss"),
    ("orders",     "userId",         "ObjectId → user","Required"),
    ("",           "name",           "String",         "Stock symbol"),
    ("",           "qty",            "Number",         "Order quantity"),
    ("",           "mode",           "Enum",           "BUY | SELL"),
    ("",           "orderType",      "Enum",           "MARKET | LIMIT (default: MARKET)"),
    ("",           "status",         "Enum",           "PENDING | EXECUTED | CANCELLED"),
    ("",           "limitPrice",     "Number",         "Nullable; user's target price"),
    ("",           "executedPrice",  "Number",         "Nullable; actual fill price"),
    ("",           "createdAt",      "Date",           "Auto (Mongoose timestamps)"),
]

st = doc.add_table(rows=len(schema_tbl_data), cols=4)
st.style = "Table Grid"
st.alignment = WD_TABLE_ALIGNMENT.CENTER
col_w = [Cm(2.5), Cm(3.5), Cm(3.2), Cm(5.0)]
for i, row_data in enumerate(schema_tbl_data):
    row = st.rows[i]
    for j, val in enumerate(row_data):
        row.cells[j].width = col_w[j]
        r = row.cells[j].paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if i == 0:
            r.bold = True
            set_cell_bg(row.cells[j], "1A5676")
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        elif val and i > 0 and j == 0:
            r.bold = True
            set_cell_bg(row.cells[j], "D0E8F5")
        elif i % 2 == 0 and j > 0:
            set_cell_bg(row.cells[j], "EAF4FB")

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 7.  MODULE DESCRIPTION
# ─────────────────────────────────────────────────────────────────────────────
heading("7. Module Description", level=1)
horizontal_rule()

modules = [
    (
        "7.1  Authentication Module",
        [
            "Handles user registration and login with PBKDF2-SHA-512 password hashing (1 000 iterations, 64-byte key).",
            "Generates a unique 32-byte hex bearer token on every login, stored in the users collection.",
            "Middleware (requireAuth) validates the token on every protected route by matching it against the database.",
            "Logout endpoint clears the stored token, invalidating the session immediately.",
            "On first login after signup, the module seeds the account with thirteen default holdings and two default positions.",
        ],
    ),
    (
        "7.2  Live Price & Watchlist Module",
        [
            "A server-side price cache stores live quotes for fourteen NSE symbols fetched via yahoo-finance2 with the .NS suffix.",
            "The cache refreshes every 30 seconds via a setInterval background job; NIFTY 50 (^NSEI) and SENSEX (^BSESN) indices are fetched simultaneously.",
            "If the Yahoo Finance API call fails, the server falls back to hardcoded demo prices so the UI never breaks.",
            "The dashboard's WatchList component polls /livePrices every 30 seconds and renders each stock's price, daily % change, and a trend arrow.",
            "An analytics modal within the WatchList shows a 7-day simulated line chart, 52-week range, and day range using Chart.js.",
        ],
    ),
    (
        "7.3  Order Management Module",
        [
            "Market Orders: On receipt of a POST /newOrder request, the server reads the live price from the cache, validates the user's cash balance (BUY) or holdings quantity (SELL), updates holdings and positions atomically, and records an EXECUTED order.",
            "Limit Orders: The order is stored as PENDING. A background job (running every 30 seconds, co-located with the price refresh) iterates all PENDING orders, checks the current cache price, and executes orders whose trigger price is met.",
            "Validation: BUY limit prices must be at most 15 % below market; SELL limit prices must be at most 15 % above market—preventing unrealistic orders.",
            "Cancel Orders: Users may cancel any PENDING limit order via DELETE /cancelOrder/:id.",
            "The Orders component in the dashboard polls /allOrders every 15 seconds and detects newly EXECUTED limit orders, surfacing a success toast notification.",
        ],
    ),
    (
        "7.4  Portfolio & Holdings Module",
        [
            "Holdings represent long-term equity positions (product type CNC). Each record stores the symbol, quantity, average buy price, and latest price.",
            "When a buy order executes, the module creates or updates the holding using a weighted-average price calculation.",
            "When a sell order executes, the quantity is decremented; if it reaches zero the holding document is deleted.",
            "Portfolio summary is computed server-side: total investment = Σ(avg × qty), current value = Σ(price × qty), P&L = current − investment.",
            "The Holdings page renders a responsive table and a bar chart (Chart.js) of current holding values.",
        ],
    ),
    (
        "7.5  Funds & Payment Simulation Module",
        [
            "Virtual funds (Paper mode): Users can add up to ₹1,00,000 per transaction via POST /addFunds. Quick-select buttons (₹5K, ₹10K, ₹25K, ₹50K) reduce friction.",
            "Live funds (Live mode): A PaymentModal simulates Card, UPI, and Net Banking flows. A processing spinner transitions to a success screen; POST /addRealFunds credits the realBalance. Maximum ₹5,00,000 per deposit.",
            "Portfolio reset: A two-step confirmation clears all holdings, positions, and orders, and resets balances to ₹50,000 via POST /resetPortfolio.",
        ],
    ),
    (
        "7.6  Dashboard UI Module",
        [
            "A two-column layout renders the WatchList in a fixed left sidebar and routes (Summary, Holdings, Orders, Funds) in the main content area.",
            "Context providers (AuthContext, ToastContext, TradingModeContext, GeneralContext) share state across the component tree without prop-drilling.",
            "The BuyActionWindow modal is activated from the WatchList or Holdings table; it refreshes the market price every 10 seconds and validates limit-price constraints before submission.",
            "Toast notifications (auto-dismiss after 4 seconds) confirm order execution, limit-order triggers, fund deposits, and error conditions.",
            "The trading-mode toggle (Paper / Live) persists to localStorage and controls which balance is displayed in the TopBar and Funds page.",
        ],
    ),
    (
        "7.7  Marketing Landing Site Module",
        [
            "A React single-page application served on Port 3000 presents six route-based pages: Home, Signup, About, Products, Pricing, and Support.",
            "The Navbar and Footer are shared across all pages. The OpenAccount call-to-action section is reused on Home, About, and other pages.",
            "The Signup page links directly to the dashboard (Port 3001), bridging the marketing site and the trading application.",
            "Bootstrap 5 provides the responsive grid; custom CSS overrides apply brand colours and typography.",
            "The Support page includes a Create Ticket form for simulated customer support submissions.",
        ],
    ),
]

for title, points in modules:
    heading(title, level=2)
    for pt in points:
        bullet(pt)
    doc.add_paragraph()

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 8.  IMPLEMENTATION DETAILS
# ─────────────────────────────────────────────────────────────────────────────
heading("8. Implementation Details", level=1)
horizontal_rule()

heading("8.1  Overview of Features Developed", level=2)

features = [
    ("User Registration & Login",            "Secure PBKDF2 password hashing, bearer-token sessions, auto-portfolio seeding"),
    ("Live Stock Watchlist",                  "14 NSE symbols + NIFTY/SENSEX via Yahoo Finance, 30-second refresh, demo fallback"),
    ("Market Order Execution",               "Instant fill at cached live price; balance and quantity validation"),
    ("Limit Order Engine",                   "Background price monitor; BUY ≤ limit, SELL ≥ limit; 15 % guard rails"),
    ("Order Cancellation",                   "Cancel PENDING limit orders; status set to CANCELLED"),
    ("Holdings Management",                  "Weighted-average price update; holdings deleted when qty reaches zero"),
    ("Positions Tracking",                   "CNC position records with P&L and isLoss flag"),
    ("Portfolio Summary",                    "Server-computed P&L, investment, current value, margin metrics"),
    ("Fund Management",                      "Virtual add-funds (₹1 L cap), simulated live payment gateway (₹5 L cap)"),
    ("Portfolio Reset",                      "Two-step confirmation, clears all data, reseeds defaults"),
    ("Analytics Charts",                     "Chart.js Doughnut (allocation), Bar (holding values), Line (7-day trend)"),
    ("Toast Notification System",            "4-second auto-dismiss, four severity levels, rendered via context"),
    ("Trading Mode Toggle",                  "Paper vs Live; persisted to localStorage; separate balances"),
    ("Responsive Multi-page Landing Site",   "Six routes, Bootstrap 5, shared Navbar/Footer, OpenAccount CTA"),
    ("Authentication Guards",                "Dashboard renders AuthPage if session invalid; routes protected client-side"),
]

tbl = doc.add_table(rows=len(features)+1, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, hdr in enumerate(["Feature", "Description"]):
    c = tbl.rows[0].cells[j]
    r = c.paragraphs[0].add_run(hdr)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    set_cell_bg(c, "1A5676")
for i, (feat, desc) in enumerate(features):
    row = tbl.rows[i+1]
    row.cells[0].paragraphs[0].add_run(feat).font.size = Pt(10)
    row.cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, "EAF4FB")

doc.add_paragraph()
heading("8.2  Code Structure (Folder Breakdown)", level=2)

folder_tree = """\
Zenvest/
├── backend/
│   ├── index.js              — Express server, all route definitions, background jobs
│   ├── package.json
│   ├── .env                  — MONGO_URL, PORT
│   ├── schemas/
│   │   ├── UserSchema.js     — User document schema
│   │   ├── HoldingsSchema.js — Holdings document schema
│   │   ├── PositionsSchema.js
│   │   └── OrdersSchema.js   — Order schema with timestamps
│   ├── model/
│   │   ├── UserModel.js      — Mongoose model export
│   │   ├── HoldingsModel.js
│   │   ├── PositionsModel.js
│   │   └── OrdersModel.js
│   ├── utils/
│   │   ├── auth.js           — PBKDF2 hashing, token generation, requireAuth middleware
│   │   └── portfolio.js      — P&L helpers, default seed data, summary calculator
│   └── seed.js               — Optional standalone seed script
│
├── dashboard/                (React trading app — Port 3001)
│   └── src/
│       ├── index.js          — Provider stack, Router, root render
│       ├── api.js            — Axios instance with auth interceptor
│       └── components/
│           ├── AuthPage.js   — Login / Signup forms
│           ├── AuthContext.js — Session state and token management
│           ├── ToastContext.js — Global toast notification system
│           ├── TradingModeContext.js — Paper/Live mode state
│           ├── GeneralContext.js — Buy/Sell window state
│           ├── Home.js       — Auth guard, renders Dashboard or AuthPage
│           ├── Dashboard.js  — Two-column layout, sidebar, content routes
│           ├── TopBar.js     — Navigation, mode toggle, logout
│           ├── WatchList.js  — Live prices, analytics modal, buy/sell triggers
│           ├── BuyActionWindow.js — Order entry modal
│           ├── Summary.js    — Portfolio overview cards
│           ├── Holdings.js   — Holdings table + bar chart
│           ├── Positions.js  — Positions table
│           ├── Orders.js     — Order history, cancel, limit-order polling
│           ├── Funds.js      — Fund management, payment modal
│           ├── DoughnoutChart.js — Holdings allocation doughnut
│           └── VerticalGraph.js  — Bar chart component
│
└── frontend/                 (React marketing site — Port 3000)
    └── src/
        ├── index.js
        └── landing_page/
            ├── Navbar.js / Footer.js / OpenAccount.js
            ├── home/         Hero, Stats, Awards, Pricing, Education
            ├── signup/       Signup.js
            ├── about/        AboutPage, Hero, Team
            ├── products/     ProductsPage, Universe, LeftSection, RightSection
            ├── pricing/      PricingPage, Brokerage
            └── support/      SupportPage, CreateTicket
"""

p = doc.add_paragraph()
r = p.add_run(folder_tree)
r.font.name = "Courier New"
r.font.size  = Pt(8)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)

heading("8.3  Key Logic Snippets", level=2)

snippets = [
    (
        "Password Hashing (backend/utils/auth.js)",
        'function hashPassword(password) {\n'
        '  const salt = crypto.randomBytes(16).toString("hex");\n'
        '  const hash = crypto.pbkdf2Sync(password, salt, 1000, 64, "sha512")\n'
        '                       .toString("hex");\n'
        '  return { salt, hash };\n'
        '}\n'
        '\n'
        'function verifyPassword(password, salt, storedHash) {\n'
        '  const hash = crypto.pbkdf2Sync(password, salt, 1000, 64, "sha512")\n'
        '                       .toString("hex");\n'
        '  return hash === storedHash;\n'
        '}',
    ),
    (
        "Market Order Execution (backend/index.js – POST /newOrder excerpt)",
        'const livePrice = priceCache[name]?.price;\n'
        'if (mode === "BUY") {\n'
        '  const totalCost = qty * livePrice;\n'
        '  if (user.cashBalance < totalCost) return res.status(400).json(...);\n'
        '  user.cashBalance -= totalCost;\n'
        '  // Upsert holding with weighted-average price\n'
        '  const existing = await HoldingsModel.findOne({ userId, name });\n'
        '  if (existing) {\n'
        '    const newAvg = (existing.avg * existing.qty + livePrice * qty)\n'
        '                   / (existing.qty + qty);\n'
        '    existing.avg = newAvg; existing.qty += qty;\n'
        '    await existing.save();\n'
        '  } else {\n'
        '    await HoldingsModel.create({ userId, name, qty, avg: livePrice, ... });\n'
        '  }\n'
        '}',
    ),
    (
        "Limit Order Background Job (backend/index.js)",
        'setInterval(async () => {\n'
        '  await refreshPriceCache();\n'
        '  const pending = await OrdersModel.find({ status: "PENDING" });\n'
        '  for (const order of pending) {\n'
        '    const market = priceCache[order.name]?.price;\n'
        '    if (!market) continue;\n'
        '    const shouldExecute =\n'
        '      (order.mode === "BUY"  && market <= order.limitPrice) ||\n'
        '      (order.mode === "SELL" && market >= order.limitPrice);\n'
        '    if (shouldExecute) await executeOrder(order, market);\n'
        '  }\n'
        '}, 30_000);',
    ),
    (
        "Axios Auth Interceptor (dashboard/src/api.js)",
        'import axios from "axios";\n'
        'const api = axios.create({ baseURL: "http://localhost:3002" });\n'
        'api.interceptors.request.use(config => {\n'
        '  const token = localStorage.getItem("zenvest_token");\n'
        '  if (token) config.headers.Authorization = `Bearer ${token}`;\n'
        '  return config;\n'
        '});\n'
        'export default api;',
    ),
]

for title, code in snippets:
    body(title, bold=True, size=10)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(code)
    r.font.name = "Courier New"
    r.font.size  = Pt(8.5)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 9.  DEPLOYMENT DETAILS
# ─────────────────────────────────────────────────────────────────────────────
heading("9. Deployment Details", level=1)
horizontal_rule()

heading("9.1  GitHub Repository", level=2)
body("Repository URL:  [Insert GitHub repository URL here]")
body("Branch:  main")
body(
    "The repository contains three top-level directories (backend, dashboard, frontend) "
    "each with their own package.json. A root-level .gitignore excludes node_modules, "
    "build artefacts, and .env files. Sensitive credentials are documented in "
    "backend/.env.example."
)

heading("9.2  Local Development Setup", level=2)
for step in [
    "Clone the repository:  git clone <repo-url>",
    "Install dependencies in each directory:  cd backend && npm install  (repeat for dashboard and frontend)",
    "Create backend/.env from backend/.env.example and set MONGO_URL and PORT",
    "Start all three services:  npm start  in each directory (backend on :3002, dashboard on :3001, frontend on :3000)",
    "Open http://localhost:3000 for the landing site; http://localhost:3001 for the trading dashboard",
]:
    bullet(step)

heading("9.3  Production Deployment (Recommended)", level=2)
deploy_data = [
    ("Component",  "Platform",         "Build Command",  "Notes"),
    ("Frontend",   "Netlify / Vercel",  "npm run build",  "Static React build; set REACT_APP_* env vars"),
    ("Dashboard",  "Netlify / Vercel",  "npm run build",  "Set REACT_APP_API_URL to backend URL"),
    ("Backend",    "Render / Railway",  "npm start",      "Set MONGO_URL, PORT env vars; enable CORS for deployed origins"),
    ("Database",   "MongoDB Atlas",     "—",              "Free tier M0 sufficient for development"),
]

dt = doc.add_table(rows=len(deploy_data), cols=4)
dt.style = "Table Grid"
dt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(deploy_data):
    for j, val in enumerate(row_data):
        c = dt.rows[i].cells[j]
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(10)
        if i == 0:
            r.bold = True
            set_cell_bg(c, "1A5676")
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        elif i % 2 == 0:
            set_cell_bg(c, "EAF4FB")

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 10. PROJECT IMPLEMENTATION STATUS TABLE
# ─────────────────────────────────────────────────────────────────────────────
heading("10. Project Implementation Status Table", level=1)
horizontal_rule()

status_data = [
    ("Module / Feature",                    "Planned", "Implemented", "Tested", "Status"),
    ("User Registration",                   "Yes",     "Yes",         "Yes",    "Complete"),
    ("User Login / Logout",                 "Yes",     "Yes",         "Yes",    "Complete"),
    ("Bearer Token Authentication",         "Yes",     "Yes",         "Yes",    "Complete"),
    ("Password Security (PBKDF2)",          "Yes",     "Yes",         "Yes",    "Complete"),
    ("Live Price Fetching (Yahoo Finance)", "Yes",     "Yes",         "Yes",    "Complete"),
    ("30-second Price Cache",               "Yes",     "Yes",         "Yes",    "Complete"),
    ("Demo Price Fallback",                 "Yes",     "Yes",         "Yes",    "Complete"),
    ("Market Order – BUY",                  "Yes",     "Yes",         "Yes",    "Complete"),
    ("Market Order – SELL",                 "Yes",     "Yes",         "Yes",    "Complete"),
    ("Limit Order – BUY",                   "Yes",     "Yes",         "Yes",    "Complete"),
    ("Limit Order – SELL",                  "Yes",     "Yes",         "Yes",    "Complete"),
    ("Background Limit Order Matcher",      "Yes",     "Yes",         "Yes",    "Complete"),
    ("Order Cancellation",                  "Yes",     "Yes",         "Yes",    "Complete"),
    ("Holdings Management",                 "Yes",     "Yes",         "Yes",    "Complete"),
    ("Positions Tracking",                  "Yes",     "Yes",         "Yes",    "Complete"),
    ("Portfolio P&L Summary",               "Yes",     "Yes",         "Yes",    "Complete"),
    ("Virtual Fund Management",             "Yes",     "Yes",         "Yes",    "Complete"),
    ("Simulated Payment Gateway",           "Yes",     "Yes",         "Yes",    "Complete"),
    ("Portfolio Reset",                     "Yes",     "Yes",         "Yes",    "Complete"),
    ("Watchlist UI with Live Prices",       "Yes",     "Yes",         "Yes",    "Complete"),
    ("Analytics Charts (Line/Bar/Doughnut)","Yes",     "Yes",         "Yes",    "Complete"),
    ("Toast Notification System",           "Yes",     "Yes",         "Yes",    "Complete"),
    ("Paper / Live Trading Mode Toggle",    "Yes",     "Yes",         "Yes",    "Complete"),
    ("NIFTY 50 / SENSEX Index Display",     "Yes",     "Yes",         "Yes",    "Complete"),
    ("Marketing Landing Site (6 pages)",    "Yes",     "Yes",         "Yes",    "Complete"),
    ("Responsive Design",                   "Yes",     "Yes",         "Yes",    "Complete"),
    ("Default Portfolio Seeding",           "Yes",     "Yes",         "Yes",    "Complete"),
]

st = doc.add_table(rows=len(status_data), cols=5)
st.style = "Table Grid"
st.alignment = WD_TABLE_ALIGNMENT.CENTER
col_widths2 = [Cm(6.5), Cm(1.8), Cm(2.5), Cm(1.8), Cm(2.2)]
for i, row_data in enumerate(status_data):
    for j, val in enumerate(row_data):
        c = st.rows[i].cells[j]
        c.width = col_widths2[j]
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if i == 0:
            r.bold = True
            set_cell_bg(c, "1A5676")
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        else:
            if j == 4 and val == "Complete":
                set_cell_bg(c, "C6EFCE")
                r.font.color.rgb = RGBColor(0x27, 0x62, 0x00)
            elif i % 2 == 0:
                set_cell_bg(c, "EAF4FB")

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 11. REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
heading("11. References", level=1)
horizontal_rule()

refs = [
    "[1]  React Documentation — https://react.dev/",
    "[2]  Node.js Documentation — https://nodejs.org/en/docs/",
    "[3]  Express.js Documentation — https://expressjs.com/",
    "[4]  Mongoose Documentation — https://mongoosejs.com/docs/",
    "[5]  MongoDB Manual — https://www.mongodb.com/docs/",
    "[6]  yahoo-finance2 npm Package — https://www.npmjs.com/package/yahoo-finance2",
    "[7]  Chart.js Documentation — https://www.chartjs.org/docs/latest/",
    "[8]  react-chartjs-2 Documentation — https://react-chartjs-2.js.org/",
    "[9]  Material-UI (MUI) Documentation — https://mui.com/",
    "[10] Bootstrap 5 Documentation — https://getbootstrap.com/docs/5.0/",
    "[11] Axios HTTP Client — https://axios-http.com/docs/intro",
    "[12] Node.js crypto Module (PBKDF2) — https://nodejs.org/api/crypto.html",
    "[13] Zerodha Kite Platform (Design Reference) — https://zerodha.com/",
    "[14] NSE India (National Stock Exchange) — https://www.nseindia.com/",
    "[15] MongoDB Atlas Cloud Database — https://www.mongodb.com/atlas",
    "[16] Netlify Deployment Platform — https://docs.netlify.com/",
    "[17] Git & GitHub — https://docs.github.com/",
    "[18] JSON Web Tokens (Bearer Tokens) — https://datatracker.ietf.org/doc/html/rfc6750",
    "[19] PBKDF2 Standard — https://datatracker.ietf.org/doc/html/rfc2898",
]

for ref in refs:
    body(ref, size=10, space_after=3)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 12. APPENDIX – SCREENSHOTS
# ─────────────────────────────────────────────────────────────────────────────
heading("12. Appendix – Screenshots", level=1)
horizontal_rule()

body(
    "The following screenshots should be captured from the running application and "
    "inserted below. Recommended screenshots are listed; replace each placeholder with "
    "the actual image.",
    italic=True
)

screenshot_list = [
    ("Figure 1",  "Marketing Landing Page – Hero Section",         "frontend  – http://localhost:3000"),
    ("Figure 2",  "Signup / Pricing Page",                         "frontend  – http://localhost:3000/signup"),
    ("Figure 3",  "Dashboard Login Screen",                        "dashboard – http://localhost:3001"),
    ("Figure 4",  "Live Watchlist with Stock Prices",              "dashboard – WatchList component"),
    ("Figure 5",  "Stock Analytics Modal (7-day chart)",           "dashboard – WatchList analytics"),
    ("Figure 6",  "Buy Order Window – Market Order",               "dashboard – BuyActionWindow"),
    ("Figure 7",  "Buy Order Window – Limit Order",                "dashboard – BuyActionWindow"),
    ("Figure 8",  "Portfolio Summary Dashboard",                   "dashboard – Summary component"),
    ("Figure 9",  "Holdings Table with Bar Chart",                 "dashboard – Holdings component"),
    ("Figure 10", "Orders Page with Status Badges",                "dashboard – Orders component"),
    ("Figure 11", "Funds Page – Virtual and Live Balances",        "dashboard – Funds component"),
    ("Figure 12", "Payment Gateway Simulation Modal",              "dashboard – PaymentModal"),
    ("Figure 13", "Toast Notification (Success)",                  "dashboard – ToastContext"),
    ("Figure 14", "NIFTY 50 / SENSEX Index Display",               "dashboard – WatchList header"),
    ("Figure 15", "Portfolio Reset Confirmation",                  "dashboard – Funds component"),
]

tbl = doc.add_table(rows=len(screenshot_list)+1, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, hdr in enumerate(["Figure", "Description", "Location / URL"]):
    c = tbl.rows[0].cells[j]
    r = c.paragraphs[0].add_run(hdr)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    set_cell_bg(c, "1A5676")
for i, (fig, desc, loc) in enumerate(screenshot_list):
    row = tbl.rows[i+1]
    row.cells[0].paragraphs[0].add_run(fig).font.size = Pt(10)
    row.cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)
    row.cells[2].paragraphs[0].add_run(loc).font.size  = Pt(9)
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, "EAF4FB")

doc.add_paragraph()
body(
    "Note: Insert actual screenshots as images into the document above. "
    "To insert: go to Insert → Pictures in Microsoft Word and select the screenshot files.",
    italic=True, size=10
)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\kmrha\OneDrive\Desktop\Zenvest\Zenvest_Project_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
